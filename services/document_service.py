import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from services.diagram_service import DiagramService

class DocumentService:
    def __init__(self):
        self.diagram_service = DiagramService()
    
    def generate_analysis_document(self, analysis_data, file_info, code_content):
        """Generate comprehensive analysis document with proper formatting"""
        try:
            # Create document
            doc = Document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add title page
            self._add_title_page(doc, analysis_data, file_info)
            
            # Add table of contents
            self._add_table_of_contents(doc)
            
            # Add executive summary
            self._add_executive_summary(doc, analysis_data)
            
            # Add code structure diagrams
            self._add_code_diagrams(doc, file_info, code_content)
            
            # Add detailed analysis
            self._add_detailed_analysis(doc, analysis_data)
            
            # Add code suggestions and recommendations
            self._add_comprehensive_suggestions(doc, analysis_data, file_info)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Code_Analysis_Report_{timestamp}.docx"
            filepath = os.path.join('temp', filename)
            
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            print(f"Error generating document: {e}")
            raise Exception(f"Document generation failed: {str(e)}")
    
    def _setup_document_styles(self, doc):
        """Setup comprehensive document styles"""
        styles = doc.styles
        
        # Main Title style
        if 'Main Title' not in [style.name for style in styles]:
            title_style = styles.add_style('Main Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        
        # Section Heading style
        if 'Section Heading' not in [style.name for style in styles]:
            section_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Calibri'
            section_style.font.size = Pt(18)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 51, 102)
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(12)
        
        # Sub Heading style
        if 'Sub Heading' not in [style.name for style in styles]:
            sub_style = styles.add_style('Sub Heading', WD_STYLE_TYPE.PARAGRAPH)
            sub_style.font.name = 'Calibri'
            sub_style.font.size = Pt(14)
            sub_style.font.bold = True
            sub_style.font.color.rgb = RGBColor(51, 51, 51)
            sub_style.paragraph_format.space_before = Pt(12)
            sub_style.paragraph_format.space_after = Pt(6)
        
        # Code Suggestion style
        if 'Code Suggestion' not in [style.name for style in styles]:
            code_sugg_style = styles.add_style('Code Suggestion', WD_STYLE_TYPE.PARAGRAPH)
            code_sugg_style.font.name = 'Consolas'
            code_sugg_style.font.size = Pt(10)
            code_sugg_style.paragraph_format.left_indent = Inches(0.5)
            code_sugg_style.paragraph_format.space_after = Pt(6)
            
            # Add background color
            shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
            code_sugg_style._element.get_or_add_pPr().append(shading_elm)
        
        # Recommendation style
        if 'Recommendation' not in [style.name for style in styles]:
            rec_style = styles.add_style('Recommendation', WD_STYLE_TYPE.PARAGRAPH)
            rec_style.font.name = 'Calibri'
            rec_style.font.size = Pt(11)
            rec_style.paragraph_format.left_indent = Inches(0.25)
            rec_style.paragraph_format.space_after = Pt(6)
    
    def _add_title_page(self, doc, analysis_data, file_info):
        """Add professional title page"""
        # Main title
        title = doc.add_paragraph('CODE ANALYSIS REPORT', style='Main Title')
        
        # Subtitle
        subtitle = doc.add_paragraph('Comprehensive Code Quality Assessment & Recommendations')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.italic = True
        subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
        
        # Add spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Report metadata table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(4)
        
        # Populate table
        metadata = [
            ('Generated On:', datetime.now().strftime('%B %d, %Y at %I:%M %p')),
            ('Quality Score:', f"{analysis_data.get('quality_score', 'N/A')}/100"),
            ('Files Analyzed:', str(len(file_info))),
            ('Analysis Type:', 'AI-Powered Comprehensive Assessment'),
            ('Report Version:', '1.0'),
            ('Generated By:', 'Test Case Generator AI System')
        ]
        
        for i, (key, value) in enumerate(metadata):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
            
            # Format cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                if cell == row.cells[0]:  # Key cell
                    cell.paragraphs[0].runs[0].bold = True
        
        # Center the table
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc):
        """Add table of contents"""
        doc.add_paragraph('TABLE OF CONTENTS', style='Section Heading')
        doc.add_paragraph()
        
        # TOC entries
        toc_entries = [
            ('1. Executive Summary', '3'),
            ('2. Code Structure Diagrams', '4'),
            ('3. Detailed Code Analysis', '6'),
            ('   3.1 Code Quality Assessment', '6'),
            ('   3.2 Complexity Analysis', '7'),
            ('   3.3 Test Coverage Gaps', '8'),
            ('   3.4 Potential Issues & Bugs', '9'),
            ('   3.5 Security Vulnerabilities', '10'),
            ('   3.6 Performance Considerations', '11'),
            ('   3.7 Design Patterns & Architecture', '12'),
            ('   3.8 Maintainability Analysis', '13'),
            ('4. AI Code Suggestions & Recommendations', '14'),
            ('   4.1 File-Specific Suggestions', '14'),
            ('   4.2 General Recommendations', '16'),
            ('   4.3 Implementation Priorities', '17')
        ]
        
        for entry, page in toc_entries:
            p = doc.add_paragraph()
            p.add_run(entry)
            
            # Add tab stops for page numbers
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6))
            
            p.add_run('\t' + page)
            p.runs[0].font.size = Pt(12)
            p.runs[1].font.size = Pt(12)
            
            if not entry.startswith('   '):  # Main sections
                p.runs[0].bold = True
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, analysis_data):
        """Add executive summary with enhanced formatting"""
        doc.add_paragraph('1. EXECUTIVE SUMMARY', style='Section Heading')
        
        # Quality score section
        quality_score = analysis_data.get('quality_score', 0)
        
        score_para = doc.add_paragraph()
        score_para.add_run('Overall Quality Score: ').bold = True
        score_run = score_para.add_run(f'{quality_score}/100')
        score_run.bold = True
        score_run.font.size = Pt(16)
        
        # Color code the score
        if quality_score >= 80:
            score_run.font.color.rgb = RGBColor(34, 139, 34)  # Green
            score_para.add_run(' (Excellent)').font.color.rgb = RGBColor(34, 139, 34)
        elif quality_score >= 60:
            score_run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
            score_para.add_run(' (Good)').font.color.rgb = RGBColor(255, 140, 0)
        else:
            score_run.font.color.rgb = RGBColor(220, 20, 60)  # Red
            score_para.add_run(' (Needs Improvement)').font.color.rgb = RGBColor(220, 20, 60)
        
        doc.add_paragraph()
        
        # Summary content
        doc.add_paragraph('Summary:', style='Sub Heading')
        summary = analysis_data.get('summary', 'Comprehensive code analysis completed successfully.')
        doc.add_paragraph(summary)
        
        # Key findings
        doc.add_paragraph('Key Findings:', style='Sub Heading')
        
        findings = [
            f"Code quality assessment completed with score: {quality_score}/100",
            f"Analysis covered {len(analysis_data.get('sections', {}))} key areas",
            "AI-powered recommendations generated for improvement",
            "Detailed suggestions provided for each code file",
            "Security, performance, and maintainability aspects evaluated"
        ]
        
        for finding in findings:
            p = doc.add_paragraph(finding, style='List Bullet')
            p.runs[0].font.size = Pt(11)
        
        doc.add_page_break()
    
    def _add_code_diagrams(self, doc, file_info, code_content):
        """Add UML-style code structure diagrams"""
        doc.add_paragraph('2. CODE STRUCTURE DIAGRAMS', style='Section Heading')
        
        doc.add_paragraph('The following diagrams illustrate the structure and organization of the analyzed code files:')
        doc.add_paragraph()
        
        try:
            for i, file_data in enumerate(file_info[:5]):  # Limit to 5 files
                doc.add_paragraph(f'2.{i+1} {file_data["name"]}', style='Sub Heading')
                
                # Generate UML diagram
                diagram_path = self.diagram_service.generate_code_diagram(
                    file_data.get('content', ''), 
                    file_data['name']
                )
                
                if diagram_path and os.path.exists(diagram_path):
                    # Add diagram
                    doc.add_picture(diagram_path, width=Inches(6.5))
                    
                    # Add centered caption
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(f'Figure 2.{i+1}: UML-style structure diagram for {file_data["name"]}')
                    caption_run.font.italic = True
                    caption_run.font.size = Pt(10)
                    caption_run.font.color.rgb = RGBColor(102, 102, 102)
                    
                    doc.add_paragraph()
                else:
                    error_para = doc.add_paragraph('Diagram generation failed for this file.')
                    error_para.runs[0].font.color.rgb = RGBColor(220, 20, 60)
                    error_para.runs[0].italic = True
        
        except Exception as e:
            doc.add_paragraph(f'Error generating diagrams: {str(e)}')
        
        doc.add_page_break()
    
    def _add_detailed_analysis(self, doc, analysis_data):
        """Add detailed analysis with proper formatting"""
        doc.add_paragraph('3. DETAILED CODE ANALYSIS', style='Section Heading')
        
        sections = analysis_data.get('sections', {})
        section_mapping = {
            'quality_assessment': '3.1 Code Quality Assessment',
            'complexity_analysis': '3.2 Complexity Analysis',
            'coverage_gaps': '3.3 Test Coverage Gaps',
            'potential_issues': '3.4 Potential Issues & Bugs',
            'security_vulnerabilities': '3.5 Security Vulnerabilities',
            'performance_considerations': '3.6 Performance Considerations',
            'design_patterns': '3.7 Design Patterns & Architecture',
            'maintainability': '3.8 Maintainability Analysis'
        }
        
        for section_key, section_title in section_mapping.items():
            doc.add_paragraph(section_title, style='Sub Heading')
            
            content = sections.get(section_key, '')
            if content:
                # Parse and format content
                formatted_content = self._format_analysis_content(content)
                for para_content in formatted_content:
                    if para_content.startswith('•') or para_content.startswith('-'):
                        p = doc.add_paragraph(para_content[1:].strip(), style='List Bullet')
                    else:
                        p = doc.add_paragraph(para_content)
                    p.runs[0].font.size = Pt(11)
            else:
                p = doc.add_paragraph('No specific findings in this category.')
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(102, 102, 102)
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_comprehensive_suggestions(self, doc, analysis_data, file_info):
        """Add comprehensive code suggestions with proper formatting"""
        doc.add_paragraph('4. AI CODE SUGGESTIONS & RECOMMENDATIONS', style='Section Heading')
        
        # File-specific suggestions
        doc.add_paragraph('4.1 File-Specific Suggestions', style='Sub Heading')
        
        for file_data in file_info:
            doc.add_paragraph(f'Recommendations for {file_data["name"]}:', style='Sub Heading')
            
            suggestions = self._generate_enhanced_suggestions(file_data)
            
            for suggestion in suggestions:
                # Suggestion title
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"💡 {suggestion['title']}")
                title_run.bold = True
                title_run.font.color.rgb = RGBColor(0, 102, 204)
                
                # Suggestion description
                desc_para = doc.add_paragraph(suggestion['description'], style='Recommendation')
                
                # Code example if available
                if 'code_example' in suggestion:
                    doc.add_paragraph('Example Implementation:', style='Sub Heading')
                    code_para = doc.add_paragraph(suggestion['code_example'], style='Code Suggestion')
                
                doc.add_paragraph()
        
        # General recommendations
        doc.add_paragraph('4.2 General Recommendations', style='Sub Heading')
        
        general_recommendations = self._get_general_recommendations()
        
        for category, recommendations in general_recommendations.items():
            doc.add_paragraph(category, style='Sub Heading')
            
            for rec in recommendations:
                p = doc.add_paragraph()
                p.add_run(f"🎯 {rec['title']}: ").bold = True
                p.add_run(rec['description'])
                p.paragraph_format.left_indent = Inches(0.25)
        
        # Implementation priorities
        doc.add_paragraph('4.3 Implementation Priorities', style='Sub Heading')
        
        priorities = [
            ('High Priority', 'Security vulnerabilities and critical bugs', RGBColor(220, 20, 60)),
            ('Medium Priority', 'Performance optimizations and code maintainability', RGBColor(255, 140, 0)),
            ('Low Priority', 'Code style improvements and documentation', RGBColor(34, 139, 34))
        ]
        
        for priority, description, color in priorities:
            p = doc.add_paragraph()
            priority_run = p.add_run(f'{priority}: ')
            priority_run.bold = True
            priority_run.font.color.rgb = color
            p.add_run(description)
    
    def _format_analysis_content(self, content):
        """Format analysis content into readable paragraphs"""
        if not content:
            return ['No analysis content available.']
        
        # Clean and split content
        cleaned_content = content.replace('**', '').replace('*', '').strip()
        paragraphs = [p.strip() for p in cleaned_content.split('\n') if p.strip()]
        
        return paragraphs[:10]  # Limit to 10 paragraphs
    
    def _generate_enhanced_suggestions(self, file_data):
        """Generate enhanced suggestions with code examples"""
        content = file_data.get('content', '')
        filename = file_data['name']
        suggestions = []
        
        if filename.endswith('.java'):
            suggestions.extend(self._analyze_java_file_enhanced(content, filename))
        elif filename.endswith('.py'):
            suggestions.extend(self._analyze_python_file_enhanced(content, filename))
        
        return suggestions
    
    def _analyze_java_file_enhanced(self, content, filename):
        """Enhanced Java file analysis with code examples"""
        suggestions = []
        
        if 'System.out.println' in content:
            suggestions.append({
                'title': 'Replace System.out with Logging Framework',
                'description': 'Replace System.out.println statements with a proper logging framework like SLF4J for better production readiness.',
                'code_example': '''// Instead of:
System.out.println("Debug message");

// Use:
private static final Logger logger = LoggerFactory.getLogger(ClassName.class);
logger.info("Debug message");'''
            })
        
        if 'public static void main' in content and content.count('\n') > 20:
            suggestions.append({
                'title': 'Extract Business Logic from Main Method',
                'description': 'Consider extracting business logic from the main method into separate methods or classes for better testability and maintainability.'
            })
        
        return suggestions
    
    def _analyze_python_file_enhanced(self, content, filename):
        """Enhanced Python file analysis with code examples"""
        suggestions = []
        
        if 'print(' in content and content.count('print(') > 2:
            suggestions.append({
                'title': 'Implement Proper Logging',
                'description': 'Replace print statements with the logging module for better control over output in production.',
                'code_example': '''# Instead of:
print("Debug message")

# Use:
import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
logger.info("Debug message")'''
            })
        
        return suggestions
    
    def _get_general_recommendations(self):
        """Get categorized general recommendations"""
        return {
            'Testing & Quality Assurance': [
                {'title': 'Unit Test Coverage', 'description': 'Aim for at least 80% code coverage with comprehensive unit tests covering edge cases and error scenarios.'},
                {'title': 'Integration Testing', 'description': 'Implement integration tests to verify component interactions and system behavior.'},
                {'title': 'Automated Testing', 'description': 'Set up continuous integration with automated test execution on code changes.'}
            ],
            'Security & Performance': [
                {'title': 'Input Validation', 'description': 'Implement comprehensive input validation and sanitization to prevent security vulnerabilities.'},
                {'title': 'Error Handling', 'description': 'Add proper error handling and logging to improve system reliability and debugging.'},
                {'title': 'Performance Monitoring', 'description': 'Implement performance monitoring and profiling to identify and address bottlenecks.'}
            ],
            'Code Maintainability': [
                {'title': 'Documentation', 'description': 'Add comprehensive code documentation including API documentation and inline comments.'},
                {'title': 'Code Organization', 'description': 'Follow SOLID principles and organize code into logical modules and packages.'},
                {'title': 'Refactoring', 'description': 'Regularly refactor code to reduce complexity and improve readability.'}
            ]
        }